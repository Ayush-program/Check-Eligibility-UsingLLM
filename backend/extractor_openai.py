import pdfplumber
import os
import re
import json
import logging
from typing import Dict, Any, List, Optional
from dotenv import load_dotenv
from groq import Groq

# Load GROQ_API_KEY from your .env file
load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ============================================================
# SPECIAL KEYWORDS — detected via regex in full PDF text
# ============================================================
SPECIAL_KEYWORDS = {
    "Malicious Code Certificate (MCC)": [
        r"malicious[\s\-]*code[\s\-]*certificate", r"\bMCC\b"
    ],
    "Non Disclosure Declaration (NDD)": [
        r"non[\s\-]*disclosure[\s\-]*declaration", r"\bNDD\b"
    ],
    "MII Compliance (Make in India Compliance)": [
        r"MII[\s\-]*compliance", r"make[\s\-]*in[\s\-]*india[\s\-]*compliance"
    ],
    "Relaxation": [r"\brelaxation\b"],
    "International Organization for Standardization (ISO)": [
        r"international[\s\-]*organization[\s\-]*for[\s\-]*standardization",
        r"\bISO\b", r"ISO[\s\-]?\d+"
    ],
    "Bureau of Indian Standards (BIS)": [
        r"bureau[\s\-]*of[\s\-]*indian[\s\-]*standards?", r"\bBIS\b"
    ],
    "Mandate":        [r"\bmandate\b"],
    "NET WORTH":      [r"net[\s\-]*worth"],
    "Escalation":     [r"\bescalation\b"],
    "Matrix":         [r"\bmatrix\b"],
    "Exemption":      [r"\bexemption\b"],
    "meeting, meet":  [r"\bmeeting\b", r"\bmeet\b"],
    "Pre Bid Detail(s)": [r"pre[\s\-]*bid[\s\-]*detail(?:s)?"],
    "Judicial / Stamp Paper": [r"\bjudicial\b", r"stamp[\s\-]*paper"],
    "Notarised / Notarized": [r"\bnotarised\b", r"\bnotarized\b"],
    "Pre Bid Queries": [r"pre[\s\-]*bid[\s\-]*queries"],
    "Exempted": [r"\bexempted\b"],
    "Blacklisted / Debarred": [
        r"black[\s\-]*listing", r"black[\s\-]*listed", r"\bblacklisting\b", 
        r"\bblacklisted\b", r"\bdebarred\b"
    ],
    "Fraud / Criminal": [r"\bcriminal\b", r"\bfraudulent\b", r"\bfraud\b"],
    "Liquidation / Court Receivership": [r"\bliquidation\b", r"court[\s\-]*receivership"]
}

# Keywords used to identify which pages to send to the AI (Bilingual)
HIGH_SIGNAL_KEYWORDS = [
    "document required from seller", "विक्रेता से मांगे गए दस्तावेज़", 
    "atc", "undertaking", "affidavit", "certificate", "compliance", 
    "turnover", "mii", "mse", "experience criteria", "past performance","submit","upload",
    "required","bidder submission",
]

class CompleteTenderExtractor:
    """
    High-Accuracy GeM Extractor.
    Filters 80-page PDFs down to 10-15 critical pages before AI analysis.
    """

    def __init__(self):
        self.api_key = os.getenv("GROQ_API_KEY")
        if not self.api_key:
            logger.error("❌ GROQ_API_KEY missing!")
        self.client = Groq(api_key=self.api_key) if self.api_key else None

    def extract_fields(self, pdf_path: str) -> Dict[str, Any]:
        """Main entry point for bid document extraction"""

        # 1. Page Filtering — first 5 pages always included (score +100 bonus)
        relevant_pages = self._get_relevant_context(pdf_path)
        raw_full_context = "\n".join([p['raw_text'] for p in relevant_pages])
        full_context = "\n".join([p['text'] for p in relevant_pages])

        if not raw_full_context:
            return self._default_data()

        # 2. Extract all basic fields — classification is read from raw text via regex
        data = self._extract_basic_info(raw_full_context)

        # === DEBUG DUMP: write raw text to file so we can see what the PDF produces ===
        try:
            import os
            debug_path = os.path.join(os.path.dirname(pdf_path), "_debug_raw_text.txt")
            with open(debug_path, "w", encoding="utf-8") as f:
                f.write(raw_full_context[:5000])
            logger.info(f"[DEBUG] Raw text dumped to: {debug_path}")
        except Exception as _de:
            logger.warning(f"[DEBUG] Could not dump raw text: {_de}")
        # === END DEBUG DUMP ===

        logger.info(f"[CLASSIFY] Detected classification_level = '{data['classification_level']}'")

        # 3. SHORT-CIRCUIT Q1 immediately — no AI, no ATC, no tokens wasted
        if data.get("classification_level") == "Q1":
            logger.info("[CLASSIFY] Q1 tender -> returning error result immediately.")
            return data
            
        data["keyword_checks"] = self._check_keywords(raw_full_context)

        # 3. Deep AI Reasoning (MAIN TENDER DOCUMENT)
        if self.client:
            try:
                logger.info(f"🤖 Analyzing {len(relevant_pages)} relevant pages via Llama 3.3 70B (Main PDF)...")
                ai_result = self._deep_ai_extraction(full_context, is_atc=False)
                data["documents_by_section"] = ai_result.get("documents_by_section", data["documents_by_section"])
            except Exception as e:
                logger.warning(f"⚠️ Main AI Pass failed: {e}")
                data = self._regex_fallback(full_context, data)

        # 4. ATC Document Processing (SEPARATE INTELLIGENT AI PASS)
        atc_link = self._extract_atc_link(pdf_path)
        data["atc_document_link"] = atc_link
        if atc_link:
            atc_text = self._fetch_atc_text(atc_link)
            if atc_text:
                logger.info("📎 Executing secondary, dedicated AI pass specifically for the ATC document.")
                # Merge any fast regex keyword flags found explicitly in the ATC
                atc_kw = self._check_keywords(atc_text)
                for k, v in atc_kw.items():
                    if v: data["keyword_checks"][k] = True
                
                # Dedicated Llama Call specifically tailored to ONLY read the ATC
                if self.client:
                    try:
                        atc_result = self._deep_ai_extraction(atc_text, is_atc=True)
                        atc_docs = atc_result.get("documents_by_section", {})
                        
                        # Cleanly merge all discovered items dynamically into the main dictionary
                        for section, docs in atc_docs.items():
                            if not docs: continue
                            # If the ATC pass puts things in "main", funnel it to "atc" to avoid confusing with the real main table
                            sec_key = section if section.lower() != "main" else "atc"
                            if sec_key not in data["documents_by_section"]:
                                data["documents_by_section"][sec_key] = []
                            data["documents_by_section"][sec_key].extend(docs)
                    except Exception as e:
                        logger.warning(f"⚠️ Secondary ATC AI Pass failed: {e}")
        
        return data

    def _preprocess_text(self, text: str) -> str:
        """Aggressive preprocessing to reduce wasteful tokens before sending to LLM."""
        if not text: return ""
        # Remove multiple spaces/tabs
        text = re.sub(r'[ \t]+', ' ', text)
        # Remove common PDF noisy footers/headers
        text = re.sub(r'Page \d+ of \d+', '', text, flags=re.IGNORECASE)
        # Remove massive alphanumeric hashes (like digital signatures)
        text = re.sub(r'\b[A-Za-z0-9]{30,}\b', '', text)
        
        # AGGRESSIVE SEMANTIC FILTERING: Drop sentences/lines that have nothing to do with documents
        keywords = {
            "document", "upload", "submit", "certificate", "affidavit", 
            "undertaking", "compliance", "eligibility", "experience", 
            "turnover", "performance", "annexure", "mandatory", "criteria", 
            "supporting", "proof", "report", "sheet", "form", "license", 
            "registration", "copy", "tender", "bid", "authorization",
            "declaration", "pact", "statement", "details", "record"
        }
        
        filtered_lines = []
        for line in text.split('\n'):
            line = line.strip()
            # Always keep very short headers or section markers (like "--- PAGE 1 ---")
            if len(line) < 15 or line.startswith("---"):
                filtered_lines.append(line)
                continue
                
            # Keep line only if it contains at least one keyword
            line_lower = line.lower()
            if any(kw in line_lower for kw in keywords):
                filtered_lines.append(line)
                
        text = "\n".join(filtered_lines)
        # Remove multiple newlines
        text = re.sub(r'\n{2,}', '\n', text)
        
        return text.strip()

    def _get_classification_from_pdf(self, pdf_path: str) -> str:
        """
        Scans EVERY page of the PDF (unfiltered, raw text) to detect Q1/Q2/Q3/Q4.
        This MUST be called before any page-scoring or semantic filtering.
        Returns the classification string e.g. 'Q1', 'Q2', 'UNIVERSAL'.
        """
        try:
            with pdfplumber.open(pdf_path) as pdf:
                all_text = ""
                # Read every page raw - no scoring, no filtering
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    all_text += t + "\n"

            logger.info(f"[CLASSIFY] Scanned {len(pdf.pages)} pages. Text sample: {repr(all_text[:300])}")

            # Pattern 1: explicit label before Q1-Q4
            m = re.search(
                r'(?:Item\s+Categ|Bid\s+Categ|Quadrant|Classification|Bid\s+Type|Bidding\s+Type)'
                r'[^\n]{0,80}?\b(Q[1-4])\b',
                all_text, re.IGNORECASE
            )
            if m:
                found = m.group(1).upper()
                logger.info(f"[CLASSIFY] Pattern-1 matched: '{found}' near '{m.group(0)[:60]}'")
                return found

            # Pattern 2: line that is ONLY or STARTS WITH Q1/Q2/Q3/Q4
            for line in all_text.splitlines():
                line_s = line.strip()
                m2 = re.match(r'^(Q[1-4])\b', line_s, re.IGNORECASE)
                if m2:
                    found = m2.group(1).upper()
                    logger.info(f"[CLASSIFY] Pattern-2 matched: '{found}' on line: '{line_s[:60]}'")
                    return found

            # Pattern 3: Q1-Q4 inside parentheses — very common in GeM e.g. "(Q1)"
            m3 = re.search(r'\((Q[1-4])\)', all_text, re.IGNORECASE)
            if m3:
                found = m3.group(1).upper()
                logger.info(f"[CLASSIFY] Pattern-3 matched: '{found}' in parentheses context")
                return found

            logger.info("[CLASSIFY] No Q1-Q4 found in full PDF scan. Returning UNIVERSAL.")
            return "UNIVERSAL"

        except Exception as e:
            logger.error(f"[CLASSIFY] Error scanning PDF for classification: {e}")
            return "UNIVERSAL"

    def _get_relevant_context(self, pdf_path: str) -> List[Dict]:
        """Filters an 80-page PDF to only high-value pages."""
        scored_pages = []
        try:
            with pdfplumber.open(pdf_path) as pdf:
                total = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    
                    # Fallback to OCR if page contains no readable text (e.g., scanned PDF)
                    if not text.strip():
                        try:
                            import pytesseract
                            import os
                            tess_path = os.getenv('TESSERACT_CMD', r'C:\Program Files\Tesseract-OCR\tesseract.exe')
                            if os.path.exists(tess_path):
                                pytesseract.pytesseract.tesseract_cmd = tess_path
                            logger.info(f"Page {i+1} appears to be scanned. Running OCR...")
                            img = page.to_image(resolution=200).original
                            text = pytesseract.image_to_string(img)
                        except ImportError:
                            logger.warning("pytesseract not installed. Cannot run OCR on scanned PDF.")
                        except Exception as e:
                            logger.warning(f"OCR failed on page {i+1}: {e}")
                            
                    text_lower = text.lower()
                    p_num = i + 1
                    
                    score = 0
                    if p_num <= 5: 
                        score += 100 # Guarantee the first 5 pages for metadata
                    
                    keywords = [
                        "document", "upload", "submit", "certificate", "affidavit", 
                        "undertaking", "compliance", "eligibility", "experience", 
                        "turnover", "past performance", "buyer added", "annexure", 
                        "scope of work", "mandatory", "criteria", "supporting documents"
                    ]
                    for kw in keywords:
                        score += text_lower.count(kw)
                        
                    if score > 0:
                        scored_pages.append({"p_num": p_num, "score": score, "text": f"\n--- PAGE {p_num} ---\n{text}"})

            # Rank top pages by keyword density. We will take the top 10 most critical pages
            top_pages = sorted(scored_pages, key=lambda x: x["score"], reverse=True)[:10]
            top_pages.sort(key=lambda x: x["p_num"])
            
            selected = [{"raw_text": p["text"], "text": self._preprocess_text(p["text"])} for p in top_pages]
            logger.info(f"📄 Filtered {total} pages down to top {len(selected)} (Density Scored & Preprocessed).")
        except Exception as e:
            logger.error(f"PDF Error: {e}")
        return selected

    def _fetch_atc_text(self, link: str) -> str:
        """Downloads the ATC document (PDF/DOCX/XLSX) and extracts text for AI analysis."""
        if not link: return ""
        logger.info(f"🌐 Downloading ATC for extraction: {link}")
        import requests, tempfile, os, pdfplumber, re
        try:
            r = requests.get(link, stream=True, timeout=30)
            if r.status_code == 200:
                # determine extension
                ext = ".pdf"
                ct = r.headers.get("Content-Type", "").lower()
                cd = r.headers.get("Content-Disposition", "")
                
                if "filename=" in cd:
                    match = re.search(r'filename="?([^"]+)"?', cd)
                    if match:
                        ext = os.path.splitext(match.group(1))[1].lower()
                elif "wordprocessingml" in ct or "msword" in ct: ext = ".docx"
                elif "spreadsheetml" in ct or "ms-excel" in ct: ext = ".xlsx"
                elif "xml" in ct or "text/xml" in ct: ext = ".xml"
                
                # Further fallback based on the URL if we still think it's a PDF
                if ext == ".pdf":
                    url_ext = os.path.splitext(link.split("?")[0])[1].lower()
                    if url_ext in ['.doc', '.docx', '.xls', '.xlsx', '.zip', '.xml']:
                        ext = url_ext

                with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                    for chunk in r.iter_content(chunk_size=8192):
                        tmp.write(chunk)
                    tmp_path = tmp.name
                
                text_content = ""
                try:
                    if ext in [".docx", ".doc"]:
                        try:
                            # Try modern docx parsing first, even if it claims to be .doc!
                            from docx import Document
                            doc = Document(tmp_path)
                            for para in doc.paragraphs:
                                if para.text.strip():
                                    text_content += para.text + "\n"
                            for table in doc.tables:
                                for row in table.rows:
                                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                                    if row_text:
                                        text_content += " | ".join(row_text) + "\n"
                        except Exception:
                            # Fallback if python-docx fails. Maybe it's a real older .doc or we don't have python-docx
                            if ext == ".docx":
                                try:
                                    import zipfile, xml.etree.ElementTree as ET
                                    with zipfile.ZipFile(tmp_path) as z:
                                        if 'word/document.xml' in z.namelist():
                                            tree = ET.fromstring(z.read('word/document.xml'))
                                            for elem in tree.iter():
                                                if elem.tag.endswith('}t') and elem.text:
                                                    text_content += elem.text + " "
                                except Exception as e:
                                    logger.warning(f"Fallback docx zip extraction failed: {e}")
                            elif ext == ".doc":
                                # Attempt rough text extraction from old .doc binary
                                with open(tmp_path, 'rb') as f:
                                    raw_data = f.read()
                                
                                header = raw_data[:20].lower()
                                if b"<html" in header or b"<!doct" in header:
                                    text_str = raw_data.decode('utf-8', errors='ignore')
                                    text_content = re.sub(r'<[^>]+>', ' ', text_str)
                                elif b"{\\rtf" in header:
                                    text_str = raw_data.decode('utf-8', errors='ignore')
                                    text_content = re.sub(r'\\[a-z]+[0-9]* ?', '', text_str)
                                    text_content = re.sub(r'[{}]', '', text_content)
                                else:
                                    # Pure binary .doc extraction (UTF-16LE / ASCII strings)
                                    # Remove null bytes to reconstruct UTF-16 encoded ascii, then strip non-printables
                                    clean_data = raw_data.replace(b'\x00', b'')
                                    clean_str = clean_data.decode('ascii', errors='ignore')
                                    text_content = re.sub(r'[^\x20-\x7E\n\r\t]', ' ', clean_str)
                                    # Safely squash horizontal spaces without destroying vertical newlines (preserves checklist formatting!)
                                    text_content = re.sub(r'[ \t]{2,}', ' ', text_content)
                    elif ext == ".xml":
                        import xml.etree.ElementTree as ET
                        try:
                            tree = ET.parse(tmp_path)
                            for elem in tree.iter():
                                if elem.text and elem.text.strip():
                                    text_content += elem.text.strip() + "\n"
                        except Exception as e:
                            logger.warning(f"Failed to parse XML: {e}")
                    elif ext == ".xlsx":
                        import zipfile, xml.etree.ElementTree as ET
                        with zipfile.ZipFile(tmp_path) as z:
                            if 'xl/sharedStrings.xml' in z.namelist():
                                tree = ET.fromstring(z.read('xl/sharedStrings.xml'))
                                for elem in tree.iter():
                                    if elem.tag.endswith('}t') and elem.text:
                                        text_content += elem.text + " "
                    elif ext == ".pdf":
                        with pdfplumber.open(tmp_path) as pdf:
                            scored_atc_pages = []
                            for i, page in enumerate(pdf.pages):
                                text = page.extract_text() or ""
                                
                                # Fallback to OCR if page contains no readable text (e.g., scanned PDF)
                                if not text.strip():
                                    try:
                                        import pytesseract
                                        import os
                                        tess_path = os.getenv('TESSERACT_CMD', r'C:\Program Files\Tesseract-OCR\tesseract.exe')
                                        if os.path.exists(tess_path):
                                            pytesseract.pytesseract.tesseract_cmd = tess_path
                                        logger.info(f"ATC Page {i+1} appears to be scanned. Running OCR...")
                                        img = page.to_image(resolution=200).original
                                        text = pytesseract.image_to_string(img)
                                    except ImportError:
                                        logger.warning("pytesseract not installed. Cannot run OCR on scanned ATC PDF.")
                                    except Exception as e:
                                        logger.warning(f"OCR failed on ATC page {i+1}: {e}")
                                        
                                text_lower = text.lower()
                                p_num = i + 1
                                score = 0
                                if p_num <= 3: score += 100
                                keywords = [
                                    "upload", "submit", "provide", "certificate", 
                                    "affidavit", "undertaking", "declaration", "mandatory",
                                    "document", "annexure", "compliance", "criteria",
                                    "eligibility", "experience", "turnover"
                                ]
                                for kw in keywords:
                                    score += text_lower.count(kw)
                                if score > 0:
                                    scored_atc_pages.append({"p_num": p_num, "score": score, "text": f"\n--- ATC PAGE {p_num} ---\n{text}"})
                            
                            # Rank the ATC pages: Keep the top 10 most important technical pages
                            top_atc_pages = sorted(scored_atc_pages, key=lambda x: x["score"], reverse=True)[:10]
                            top_atc_pages.sort(key=lambda x: x["p_num"])
                            for p in top_atc_pages:
                                text_content += p["text"]
                    else:
                        logger.warning(f"ATC format {ext} is not natively supported for text parsing yet. Skipping textual merge.")
                except Exception as e:
                    logger.warning(f"ATC couldn't be fully extracted as {ext}: {e}")
                
                try:
                    os.unlink(tmp_path)
                except Exception as del_e:
                    logger.warning(f"Could not delete temp ATC file: {del_e}")
                    
                return self._preprocess_text(text_content)
        except Exception as e:
            logger.error(f"Error fetching ATC text: {e}")
            return ""

    def _deep_ai_extraction(self, text: str, is_atc: bool = False) -> Dict[str, Any]:
        """Exhaustive AI extraction - rewrites prompt per document type for maximum recall."""

        if is_atc:
            system_msg = (
                "You are a GeM tender compliance expert. Extract EVERY document a bidder must submit from this ATC document. "
                "Scan every line, numbered list, bullet, and table. Look for: certificates, annexures, forms, declarations, "
                "affidavits, pacts, undertakings, authorizations, compliance sheets, data sheets, proofs. "
                "CRITICAL RULES: "
                "1. STRICTLY ANTI-HALLUCINATORY: Do not invent or include items not explicitly stated as required documents. "
                "2. NO NON-DOCUMENTS: DO NOT extract sentences, generic rules, criteria descriptions, hardware specs, action items, or terms and conditions. ONLY exact document names. "
                "3. FULL AND CLEAR NAMING: Extract the FULL, specific name of the document EXACTLY as intended in the text to avoid any confusion. Do NOT over-simplify (e.g. if it says 'Non-Disclosure Agreement (NDA) as per Annexure V', keep the relevant descriptive name so the bidder knows exactly what to upload). DO NOT extract entire sentences or action verbs. Remove any leading bullet points or numbering (like '1.' or '*'). "
                "4. DYNAMIC CATEGORIZATION: Group documents into logical, specific categories (e.g., 'Certificates', 'Annexures', 'Declarations') instead of lumping them together. "
                "Return JSON: {\"documents_by_section\": {\"atc\": [\"...\"], \"<Dynamic Category>\": [\"...\"]}}."
            )
        else:
            system_msg = (
                "You are a GeM tender compliance expert. Extract EVERY document a bidder must submit from this tender. "
                "Scan every line, table row, and numbered list. Focus heavily on: 'Document required from seller' tables, "
                "eligibility criteria, financial proofs, technical certificates, annexures, forms, declarations. "
                "CRITICAL RULES: "
                "1. STRICTLY ANTI-HALLUCINATORY: Do not invent or include items not explicitly stated as required documents. "
                "2. NO NON-DOCUMENTS: DO NOT extract sentences, generic rules, criteria descriptions, hardware specs, action items, or terms and conditions. ONLY exact document names. "
                "3. FULL AND CLEAR NAMING: Extract the FULL, specific name of the document EXACTLY as intended in the text to avoid any confusion. Do NOT over-simplify (e.g. if it says 'Non-Disclosure Agreement (NDA) as per Annexure V', keep the relevant descriptive name so the bidder knows exactly what to upload). DO NOT extract entire sentences or action verbs. Remove any leading bullet points or numbering (like '1.' or '*'). "
                "4. DYNAMIC CATEGORIZATION: Group documents into logical, specific categories (e.g., 'Financial Documents', 'Certificates'). Use 'main' for the primary documents table. "
                "Return JSON: {\"documents_by_section\": {\"main\": [\"...\"], \"<Dynamic Category 1>\": [\"...\"]}}."
            )

        merged: Dict[str, list] = {}
        
        if not text.strip():
            return {"documents_by_section": merged}

        user_msg = f"Extract ALL required bidder documents entirely from this text:\n\n{text}"
        
        # Dual-Model Load Balancing to bypass Groq's 12,000 TPM limit
        # The main document uses the 70B model, while the ATC uses the 8B model. 
        # Groq tracks rate limits PER MODEL, allowing us to process up to 24,000+ total tokens per minute!
        model_selection = "llama-3.1-8b-instant" if is_atc else "llama-3.3-70b-versatile"
        logger.info(f"Using Model: {model_selection} for extraction.")

        try:
            response = self.client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": system_msg},
                    {"role": "user", "content": user_msg}
                ],
                temperature=0,
                max_tokens=600,
                response_format={"type": "json_object"}
            )
            result = json.loads(response.choices[0].message.content)
            chunk_sections = result.get("documents_by_section", {})

            for section_key, docs in chunk_sections.items():
                if not isinstance(docs, list):
                    continue
                if section_key not in merged:
                    merged[section_key] = []
                for doc in docs:
                    if doc and isinstance(doc, str):
                        # Clean the document name perfectly
                        clean_doc = re.sub(r'^[\s\d\.\*\-]+', '', doc)  # Remove leading numbers and bullets like "1. ", "* "
                        clean_doc = re.sub(r'[\s\.\:\;\-]+$', '', clean_doc)  # Remove trailing punctuation
                        clean_doc = clean_doc.strip('\'"').strip()
                        
                        # Filtering checks
                        if len(clean_doc) < 3: # Too short
                            continue
                        if len(clean_doc.split()) > 15: # Way too long, likely a hallucinated sentence/clause
                            continue
                            
                        if clean_doc not in merged[section_key]:
                            merged[section_key].append(clean_doc)

        except Exception as e:
            logger.error(f"AI Extraction failed: {e}")

        return {"documents_by_section": merged}




    # ============================================================
    # REGEX FALLBACK
    # ============================================================

    def _regex_fallback(self, text: str, data: Dict) -> Dict:
        """Pure regex extraction when AI is unavailable or fails"""

        sections = {
            "main": [], "atc": [], "financial": [],
            "technical": [], "certificates": [], "additional": []
        }

        # MAIN: "Document required from seller" field
        doc_match = re.search(
            r'Document required from seller[:\s/]*\n?(.*?)(?=\*In case|\n\n|\Z)',
            text, re.IGNORECASE | re.DOTALL
        )
        if doc_match:
            for item in re.split(r'[,\n]+', doc_match.group(1)):
                item = item.strip().strip('*').strip()
                if item and len(item) > 2 and len(item) < 150:
                    sections["main"].append(item)

        # ATC: look for numbered items in buyer-added ATC
        atc_match = re.search(
            r'Buyer Added Bid Specific.*?(.*?)(?=Disclaimer|अस्वीकरण|\Z)',
            text, re.IGNORECASE | re.DOTALL
        )
        if atc_match:
            numbered = re.findall(r'\d+\.\s+([A-Z][^.\n]{10,100})', atc_match.group(1))
            for item in numbered:
                if any(kw in item.lower() for kw in [
                    'upload', 'submit', 'certificate', 'affidavit', 'undertaking',
                    'declaration', 'document', 'sheet', 'proof', 'pact', 'form',
                    'brochure', 'data sheet', 'letter', 'copy', 'mandate'
                ]):
                    sections["atc"].append(item.strip()[:100])

        # FINANCIAL
        financial_patterns = [
            (r'EMD Amount', "EMD (Earnest Money Deposit)"),
            (r'Performance.*?Bank Guarantee|ePBG', "Performance Bank Guarantee (PBG)"),
            (r'Audited Balance Sheet', "Audited Balance Sheets"),
            (r'CA.*?turnover|Chartered Accountant.*?turnover', "CA Certificate for Turnover"),
            (r'Profit\s*[&]\s*Loss|Profit and Loss', "Profit & Loss Statement"),
            (r'Income Tax Acknowledgement', "Income Tax Acknowledgement"),
            (r'Net Worth', "Net Worth Certificate"),
            (r'ITR.*?\d{4}', "Income Tax Returns"),
        ]
        for pat, name in financial_patterns:
            if re.search(pat, text, re.IGNORECASE):
                sections["financial"].append(name)

        # TECHNICAL
        technical_patterns = [
            (r'Data Sheet.*?product|product.*?data sheet', "Product Data Sheet"),
            (r'catalogue|brochure', "Product Catalogue / Brochure"),
            (r'Technical.*?Compliance.*?Sheet', "Technical Compliance Sheet"),
            (r'Compliance.*?BoQ', "BoQ Compliance Document"),
            (r'Technical Bid', "Technical Bid Documents"),
        ]
        for pat, name in technical_patterns:
            if re.search(pat, text, re.IGNORECASE):
                sections["technical"].append(name)

        # CERTIFICATES
        certificate_patterns = [
            (r'ISO\s*9001', "ISO 9001 Certificate"),
            (r'\bBIS\b', "BIS Certificate"),
            (r'Drug License', "Drug License"),
            (r'Udyam Registration', "Udyam Registration Certificate"),
            (r'DPIIT.*?Startup|Startup.*?DPIIT', "DPIIT Startup Certificate"),
            (r'GST.*?Registration|GSTIN', "GST Registration Certificate"),
            (r'PAN Card', "PAN Card"),
            (r'MSME', "MSME/Udyam Certificate"),
        ]
        for pat, name in certificate_patterns:
            if re.search(pat, text, re.IGNORECASE):
                sections["certificates"].append(name)

        # ADDITIONAL
        additional_patterns = [
            (r'Land Border Declaration', "Land Border Declaration"),
            (r'Integrity Pact', "Integrity Pact"),
            (r'MII.*?certif|Make.*?India.*?certif', "MII Local Content Certificate"),
            (r'Malicious Code', "Malicious Code Certificate (MCC)"),
            (r'EFT Mandate', "EFT Mandate"),
            (r'Undertaking', "Undertaking Letter"),
            (r'Affidavit', "Affidavit"),
        ]
        for pat, name in additional_patterns:
            if re.search(pat, text, re.IGNORECASE):
                sections["additional"].append(name)

        # Deduplicate within sections
        for sec in sections:
            sections[sec] = self._deduplicate(sections[sec])

        data["documents_by_section"] = sections
        return data

    # ============================================================
    # BASIC INFO (regex, fast)
    # ============================================================

    def _extract_basic_info(self, text: str) -> Dict[str, Any]:
        data = {
            "bid_number": None,
            "bid_end_date": None,
            "item_category": None,
            "classification_level": "UNIVERSAL",
            "turnover": 0,
            "oem_turnover": 0,
            "experience": 0,
            "past_performance": 0,
            "mse_status": "No",
            "startup_status": "No",
            "emd_amount": None,
            "documents_by_section": {
                "main": [], "atc": [], "financial": [],
                "technical": [], "certificates": [], "additional": []
            },
            "keyword_checks": {},
            "atc_document_link": None
        }

        # Bid Number
        m = re.search(r'GEM/\d{4}/[A-Z]/\d+', text)
        if m:
            data["bid_number"] = m.group()

        # Bid End Date
        m = re.search(r'Bid End Date[/\s:]*(?:Time)?\s*(\d{2}-\d{2}-\d{4})', text, re.IGNORECASE)
        if m:
            data["bid_end_date"] = m.group(1)

        # Item Category
        m = re.search(r'Item Category\s*[:\n]+\s*([^\n]+)', text, re.IGNORECASE)
        if m:
            raw_cat = m.group(1).strip()
            q = re.search(r'\(Q([1-4])\)', raw_cat)
            if q:
                data["classification_level"] = f"Q{q.group(1)}"
                raw_cat = re.sub(r'\s*\(Q[1-4]\)', '', raw_cat).strip()
            data["item_category"] = raw_cat

        if data["classification_level"] == "UNIVERSAL":
            # Robust multi-pattern Q1-Q4 fallback detection
            # Pattern A: labeled keyword followed by Q1-Q4 on the same line
            cl_match = re.search(
                r'(?:Item\s+Categ|Bid\s+Categ|Quadrant|Classification|Bid\s+Type|Bidding\s+Type)'
                r'[^\n]{0,80}?\b(Q[1-4])\b',
                text, re.IGNORECASE
            )

            # Pattern B: Q1-Q4 inside parentheses — very common: "(Q1)"
            if not cl_match:
                cl_match = re.search(r'\((Q[1-4])\)', text, re.IGNORECASE)

            # Pattern C: line that STARTS with Q1/Q2/Q3/Q4 (short label lines)
            if not cl_match:
                for line in text.splitlines():
                    ls = line.strip()
                    m2 = re.match(r'^(Q[1-4])\b', ls, re.IGNORECASE)
                    if m2 and len(ls) < 10:
                        cl_match = m2
                        break

            if cl_match:
                data["classification_level"] = cl_match.group(1).upper()
                logger.info(f"[CLASSIFY] Fallback regex matched: {data['classification_level']}")

        # Turnover
        m = re.search(
            r'Minimum Average Annual Turnover.*?(\d+(?:\.\d+)?)\s*(Lakh|Crore)',
            text, re.IGNORECASE
        )
        if m:
            val, unit = float(m.group(1)), m.group(2).lower()
            data["turnover"] = int(val * (10_000_000 if unit == "crore" else 100_000))

        # Experience
        m = re.search(r'Years of Past Experience Required.*?(\d+)\s*Year', text, re.IGNORECASE)
        if m:
            data["experience"] = int(m.group(1))

        # Past Performance
        m = re.search(r'Past Performance[:\s/]*(\d+)\s*%', text, re.IGNORECASE)
        if m:
            data["past_performance"] = int(m.group(1))

        # MSE Status
        m = re.search(r'MSE (?:Exemption|Relaxation) for Years.*?(Yes|No)', text, re.IGNORECASE)
        if m:
            data["mse_status"] = "Yes" if m.group(1).lower() == "yes" else "No"

        # Startup Status
        m = re.search(r'Startup (?:Exemption|Relaxation) for Years.*?(Yes|No)', text, re.IGNORECASE)
        if m:
            data["startup_status"] = "Yes" if m.group(1).lower() == "yes" else "No"

        # EMD Amount
        m = re.search(r'EMD Amount[:\s]*(\d+)', text, re.IGNORECASE)
        if m:
            data["emd_amount"] = int(m.group(1))

        return data

    # ============================================================
    # KEYWORD CHECKS
    # ============================================================

    def _check_keywords(self, text: str) -> Dict[str, bool]:
        results = {}
        for keyword, patterns in SPECIAL_KEYWORDS.items():
            found = any(re.search(p, text, re.IGNORECASE) for p in patterns)
            results[keyword] = found
            logger.info(f"  {'✅' if found else '❌'} {keyword}")
        return results

    # ============================================================
    # ATC LINK
    # ============================================================

    def _extract_atc_link(self, pdf_path: str) -> Optional[str]:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    if re.search(r"Buyer uploaded ATC document|Click here to view", text, re.IGNORECASE):
                        if page.hyperlinks:
                            for hl in page.hyperlinks:
                                uri = hl.get("uri", "")
                                if "http" in uri:
                                    logger.info(f"🔗 ATC link: {uri}")
                                    return uri
        except Exception as e:
            logger.error(f"ATC link error: {e}")
        return None

    # ============================================================
    # UTILITIES
    # ============================================================

    def _deduplicate(self, docs: List[str]) -> List[str]:
        seen, result = set(), []
        for doc in docs:
            doc = doc.strip()
            if not doc or len(doc) < 3:
                continue
            key = re.sub(r'[^a-z0-9]', '', doc.lower())
            if key and len(key) > 3 and key not in seen:
                seen.add(key)
                result.append(doc)
        return result

    def _default_data(self) -> Dict[str, Any]:
        return {
            "bid_number": None,
            "bid_end_date": None,
            "item_category": None,
            "classification_level": "UNIVERSAL",
            "turnover": 0,
            "oem_turnover": 0,
            "experience": 0,
            "past_performance": 0,
            "mse_status": "No",
            "startup_status": "No",
            "emd_amount": None,
            "documents_by_section": {
                "main": [], "atc": [], "financial": [],
                "technical": [], "certificates": [], "additional": []
            },
            "keyword_checks": {k: False for k in SPECIAL_KEYWORDS},
            "atc_document_link": None
        }


# ============================================================
# CLI TEST
# ============================================================

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python extractor.py <path_to_pdf>")
        sys.exit(1)

    extractor = CompleteTenderExtractor()
    result = extractor.extract_fields(sys.argv[1])

    print("\n" + "=" * 80)
    print("📄 GeM TENDER EXTRACTION RESULT")
    print("=" * 80)
    
    # Print basic fields
    basic_fields = ["bid_number", "bid_end_date", "item_category", "classification_level",
                    "turnover", "experience", "past_performance", "mse_status",
                    "startup_status", "emd_amount"]
    
    for field in basic_fields:
        val = result.get(field, "N/A")
        if field == "turnover" and val:
            val = f"₹{val:,}"
        print(f"  {field:<22}: {val}")

    print(f"\n  ATC Link: {result.get('atc_document_link', 'None')}")

    # Print documents by section
    print("\n📂 DOCUMENTS BY SECTION:")
    labels = {
        "main":         "📋 Main (Document required from seller)",
        "atc":          "📄 ATC / Buyer Added Terms",
        "financial":    "💰 Financial",
        "technical":    "🔧 Technical",
        "certificates": "📜 Certificates",
        "additional":   "📎 Additional",
    }
    
    for sec, docs in result.get("documents_by_section", {}).items():
        if not docs: continue
        label = labels.get(sec, f"📌 {sec.replace('_', ' ').title()}")
        print(f"\n  {label} ({len(docs)}):")
        for d in docs:
            print(f"     • {d}")

    # Flatten and calculate total unique documents
    unique_docs = set()
    for sec_docs in result.get("documents_by_section", {}).values():
        for doc in sec_docs:
            unique_docs.add(doc.lower().strip())
    
    print(f"\n📊 TOTAL UNIQUE DOCUMENTS (across all sections): {len(unique_docs)}")

    # Print keyword checks
    print("\n🔍 KEYWORD CHECKS:")
    for keyword, found in result.get("keyword_checks", {}).items():
        print(f"  {'✅' if found else '❌'} {keyword}")

    # Save to JSON file
    with open("extraction_result.json", "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False, default=str)
    print("\n✅ Saved to extraction_result.json")